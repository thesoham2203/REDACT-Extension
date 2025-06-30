import spacy
import re
from faker import Faker
from docx import Document
from docx.shared import RGBColor

class DOCRedactor:
    def _init_(self):
        self.nlp = spacy.load("en_core_web_sm")
        self.fake = Faker()

    def extract_text_and_coordinates(self, docx_path):
        doc = Document(docx_path)
        text = ""
        paragraphs = []

        for para_num, para in enumerate(doc.paragraphs):
            text += para.text + "\n"
            paragraphs.append({"text": para.text, "para_num": para_num})

        return text, paragraphs

    def extract_sensitive_data(self, text):
        doc = self.nlp(text)
        sensitive_data = []

        for ent in doc.ents:
            if ent.label_ in ['PERSON', 'GPE', 'ORG', 'DATE', 'MONEY']:
                sensitive_data.append((ent.start, ent.end, ent.text, ent.label_))

        # Patterns for additional sensitive data
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
        phone_pattern = r'\b(\+?[0-9]{1,3})?[-. ]?(\(?\d{3}\)?[-. ]?\d{3}[-. ]?\d{4})\b'
        ipv4_pattern = r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b'
        ipv6_pattern = r'\b(?:[A-Fa-f0-9]{1,4}:){7}[A-Fa-f0-9]{1,4}\b'
        time_pattern = r'\b([01]?[0-9]|2[0-3]):([0-5][0-9])\b'  # Time pattern HH:MM

        # Find matching patterns in text
        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        ipv4s = re.findall(ipv4_pattern, text)
        ipv6s = re.findall(ipv6_pattern, text)
        times = re.findall(time_pattern, text)  # Matching time patterns

        # Append to sensitive data
        for email in emails:
            sensitive_data.append((0, 0, email, 'EMAIL'))
        for phone in phones:
            sensitive_data.append((0, 0, phone[1], 'PHONE'))
        for ipv4 in ipv4s:
            sensitive_data.append((0, 0, ipv4, 'IPV4'))
        for ipv6 in ipv6s:
            sensitive_data.append((0, 0, ipv6, 'IPV6'))
        for time in times:
            sensitive_data.append((0, 0, ":".join(time), 'TIME'))  # Adding matched times

        return sensitive_data

    def generate_synthetic_data(self, label):
        if label == 'EMAIL':
            return self.fake.email()
        elif label == 'PHONE':
            return self.fake.phone_number()
        elif label == 'PERSON':
            return self.fake.name()
        elif label == 'GPE':
            return self.fake.city()
        elif label == 'ORG':
            return self.fake.company()
        elif label == 'DATE':
            return self.fake.date()
        elif label == 'MONEY':
            return self.fake.pricetag()
        elif label == 'IPV4':
            return self.fake.ipv4()
        elif label == 'IPV6':
            return self.fake.ipv6()
        elif label == 'TIME':
            return self.fake.time()  # Generate synthetic time
        else:
            return "SYNTHETIC_DATA"

    def redact_blackout(self, para, sensitive_text):
        if sensitive_text in para.text:
            # Replace sensitive text with a black box (█ symbol)
            para.text = para.text.replace(sensitive_text, "█" * len(sensitive_text))

    def redact_blur(self, para, sensitive_text):
        if sensitive_text in para.text:
            # Replace sensitive text with dashes (for blur effect)
            para.text = para.text.replace(sensitive_text, '-' * len(sensitive_text))

    def replace_with_synthetic_data(self, para, sensitive_text, label):
        if sensitive_text in para.text:
            # Generate synthetic data using Faker
            synthetic_text = self.generate_synthetic_data(label)
            para.text = para.text.replace(sensitive_text, synthetic_text)

    def redact_table(self, table, sensitive_data, action):
        for row in table.rows:
            for cell in row.cells:
                for _, _, sensitive_text, label in sensitive_data:
                    # Ensure that sensitive data in the table cell is also redacted
                    if sensitive_text in cell.text:
                        if action == 'b':
                            self.redact_blackout(cell, sensitive_text)
                        elif action == 'x':
                            self.redact_blur(cell, sensitive_text)
                        elif action == 's':
                            self.replace_with_synthetic_data(cell, sensitive_text, label)

    def process_docx(self, docx_path, redact_level, action):
        doc_text, paragraphs = self.extract_text_and_coordinates(docx_path)
        sensitive_data = self.extract_sensitive_data(doc_text)

        if not sensitive_data:
            print("No sensitive data found.")
            return

        doc = Document(docx_path)
        entities_to_redact = set()

        # Adjust redaction levels based on sensitivity
        if redact_level >= 25:
            entities_to_redact.update(['EMAIL', 'PHONE', 'IPV4', 'IPV6'])
        if redact_level >= 50:
            entities_to_redact.update(['DATE'])
        if redact_level >= 75:
            entities_to_redact.update(['MONEY', 'ORG', 'GPE'])  # Redact city names (GPE)
        if redact_level == 100:
            entities_to_redact.update(['PERSON'])

        # Filter the sensitive data based on redact level
        filtered_data = [data for data in sensitive_data if data[3] in entities_to_redact]

        # Redact in paragraphs (main body)
        for _, _, sensitive_text, label in filtered_data:
            for para in doc.paragraphs:
                if sensitive_text in para.text:
                    if action == 'b':
                        self.redact_blackout(para, sensitive_text)
                    elif action == 'x':
                        self.redact_blur(para, sensitive_text)
                    elif action == 's':
                        self.replace_with_synthetic_data(para, sensitive_text, label)

        # Redact in tables
        for table in doc.tables:
            self.redact_table(table, filtered_data, action)

        # Redact in headers (ensure header paragraphs are processed)
        for section in doc.sections:
            for header in section.header.paragraphs:  # Corrected header access
                for _, _, sensitive_text, label in filtered_data:
                    if sensitive_text in header.text:
                        if action == 'b':
                            self.redact_blackout(header, sensitive_text)
                        elif action == 'x':
                            self.redact_blur(header, sensitive_text)
                        elif action == 's':
                            self.replace_with_synthetic_data(header, sensitive_text, label)

        redacted_docx_path = docx_path.replace('.docx', '_redacted.docx')
        doc.save(redacted_docx_path)
        return redacted_docx_path

# def main():
#     docx_path = "C:/Users/Rajesh/Downloads/SIH_Dataset/Doc/MAPITSecurityAuditInitiationDocument.docx"

#     redact_level = int(input("Enter the redaction level (0 to 100): "))
#     action = input("Enter the redaction action (b = blackout, x = blur, s = synthetic data): ")

#     redactor = DOCXRedactor()
#     redacted_docx = redactor.process_docx(docx_path, redact_level, action)
    
#     if redacted_docx:
#         print(f"Redacted DOCX saved as {redacted_docx}")

# if __name__ == "__main__":
#     main()